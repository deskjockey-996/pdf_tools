from flask import Blueprint, request, send_file, render_template, jsonify
import os
import tempfile
import zipfile
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from werkzeug.utils import secure_filename
import io
import logging

word2pdf = Blueprint('word2pdf', __name__)

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 字体配置
FONT_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'fonts')
PREFERRED_FONTS = [
    ('NotoSansSCMedium.ttf', 'NotoSansSC'),
    ('NotoSansCJKsc-Regular.otf', 'NotoSansCJK')
]
DEFAULT_FONT = 'Helvetica'

def register_chinese_font():
    """注册中文字体"""
    try:
        # 检查首选字体
        for font_file, font_name in PREFERRED_FONTS:
            font_path = os.path.join(FONT_PATH, font_file)
            if os.path.exists(font_path):
                try:
                    logger.info(f"正在注册字体: {font_path}")
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    logger.info(f"成功注册字体: {font_name}")
                    return font_name
                except Exception as e:
                    logger.error(f"注册字体 {font_file} 失败: {str(e)}")
                    continue

        # 如果首选字体都不可用，尝试其他 .ttf 或 .otf 字体
        if os.path.exists(FONT_PATH):
            font_files = [f for f in os.listdir(FONT_PATH) 
                         if f.endswith(('.ttf', '.otf')) 
                         and not f.startswith('fa-')]  # 排除 FontAwesome 字体
            
            if font_files:
                font_file = font_files[0]
                font_path = os.path.join(FONT_PATH, font_file)
                font_name = os.path.splitext(font_file)[0]
                try:
                    logger.info(f"尝试使用备选字体: {font_path}")
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    logger.info(f"成功注册备选字体: {font_name}")
                    return font_name
                except Exception as e:
                    logger.error(f"注册备选字体失败: {str(e)}")

        logger.warning("未找到可用的中文字体，将使用默认字体")
        return DEFAULT_FONT
    except Exception as e:
        logger.error(f"字体注册过程出错: {str(e)}")
        return DEFAULT_FONT

@word2pdf.route('/')
def word2pdf_page():
    return render_template('word2pdf.html')

def convert_docx_to_pdf(docx_path, pdf_path):
    """将.docx文档转换为PDF"""
    try:
        # 注册字体
        font_name = register_chinese_font()
        logger.info(f"使用字体: {font_name}")

        # 读取Word文档
        doc = Document(docx_path)
        
        # 创建PDF文档
        pdf = SimpleDocTemplate(
            pdf_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # 创建样式
        styles = getSampleStyleSheet()
        style = ParagraphStyle(
            'CustomStyle',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=12,
            leading=14,
            encoding='utf-8'
        )
        
        # 准备内容
        story = []
        
        # 处理每个段落
        for para in doc.paragraphs:
            if para.text.strip():  # 只处理非空段落
                try:
                    p = Paragraph(para.text, style)
                    story.append(p)
                    story.append(Spacer(1, 12))
                except Exception as e:
                    logger.error(f"处理段落时出错: {str(e)}")
                    # 如果处理失败，尝试使用纯文本
                    p = Paragraph(para.text.encode('utf-8').decode('utf-8', 'ignore'), style)
                    story.append(p)
                    story.append(Spacer(1, 12))
        
        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                try:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        p = Paragraph(row_text, style)
                        story.append(p)
                        story.append(Spacer(1, 12))
                except Exception as e:
                    logger.error(f"处理表格时出错: {str(e)}")
        
        # 生成PDF
        pdf.build(story)
        return True
    except Exception as e:
        logger.error(f"转换错误: {str(e)}")
        return False

@word2pdf.route('', methods=['POST'])
def convert_word_to_pdf():
    if 'wordFiles' not in request.files:
        return jsonify({'error': '未找到文件'}), 400
    
    files = request.files.getlist('wordFiles')
    if not files or files[0].filename == '':
        return jsonify({'error': '未选择文件'}), 400

    try:
        # 创建临时目录
        with tempfile.TemporaryDirectory() as temp_dir:
            # 创建ZIP文件
            zip_path = os.path.join(temp_dir, 'converted_pdfs.zip')
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                converted_count = 0
                for file in files:
                    if not file.filename.lower().endswith('.docx'):
                        continue

                    try:
                        # 保存上传的Word文件
                        word_path = os.path.join(temp_dir, secure_filename(file.filename))
                        file.save(word_path)

                        # 设置输出PDF路径
                        base_name = os.path.splitext(secure_filename(file.filename))[0]
                        pdf_filename = f"{base_name}.pdf"
                        pdf_path = os.path.join(temp_dir, pdf_filename)

                        # 转换文档
                        if convert_docx_to_pdf(word_path, pdf_path):
                            # 添加到ZIP文件
                            zipf.write(pdf_path, pdf_filename)
                            converted_count += 1
                        else:
                            logger.error(f"转换文件失败: {file.filename}")
                        
                        # 删除临时文件
                        os.remove(word_path)
                        if os.path.exists(pdf_path):
                            os.remove(pdf_path)
                    except Exception as e:
                        logger.error(f"处理文件 {file.filename} 时出错: {str(e)}")

                if converted_count == 0:
                    return jsonify({'error': '没有文件被成功转换'}), 500

            # 返回ZIP文件
            return send_file(
                zip_path,
                as_attachment=True,
                download_name='converted_pdfs.zip',
                mimetype='application/zip'
            )

    except Exception as e:
        logger.error(f"转换失败: {str(e)}")
        return jsonify({'error': f'转换失败: {str(e)}'}), 500 